from unittest.mock import Base
import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import requests
import time
import re
import os
from typing import Dict, Optional
from dotenv import load_dotenv
from utils.auth import initialize_auth_state, is_logged_in

load_dotenv()

try:
    import PyPDF2
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

st.set_page_config(
    page_title="Text Similarity Analyzer",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
    }
            
    .main-header p {
        margin-bottom: 0;
    }
    
    .metric-card {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        padding: 1rem;
        border-radius: 10px;
        border-left: 5px solid #667eea;
        margin: 0.5rem 0;
    }
    
    .similarity-high {
        background: linear-gradient(135deg, #d4edda 0%, #c3e6cb 100%);
        border-left-color: #28a745;
    }
    
    .similarity-medium {
        background: linear-gradient(135deg, #fff3cd 0%, #ffeaa7 100%);
        border-left-color: #ffc107;
    }
    
    .similarity-low {
        background: linear-gradient(135deg, #f8d7da 0%, #f1c0c7 100%);
        border-left-color: #dc3545;
    }
    
    .speed-indicator {
        padding: 0.25rem 0.5rem;
        border-radius: 12px;
        font-size: 0.8rem;
        font-weight: bold;
        margin-left: 0.5rem;
    }
    
    .speed-fast { background: #d4edda; color: #155724; }
    .speed-medium { background: #fff3cd; color: #856404; }
    .speed-slow { background: #f8d7da; color: #721c24; }
    
    .api-status {
        padding: 0.5rem 1rem;
        border-radius: 8px;
        margin: 0.5rem 0;
        border-left: 4px solid;
    }
    
    .api-connected { background: #d4edda; border-left-color: #28a745; color: #155724; }
    .api-disconnected { background: #f8d7da; border-left-color: #dc3545; color: #721c24; }
    
    .stButton > button {
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 15px rgba(102, 126, 234, 0.4);
    }
    
    div[data-testid="column"]:last-child {
        display: flex;
        align-items: flex-start;
        justify-content: center;
    }
            
    [data-testid="stHeading"] a {
        display: none !important;
    }
            
    [data-testid="stSidebarNav"] {
        display: none !important;
    }
    
            
    [data-testid="stSidebarHeader"] {
        display: none !important;
    }

    [data-testid="collapsedControl"] {
        display: none
    }       

    h1 a, h2 a, h3 a, h4 a, h5 a, h6 a {
        display: none !important;
    }
            
    button[title="View fullscreen"] {
        visibility: hidden;
    }
            
</style>
""", unsafe_allow_html=True)

hide_streamlit_style = """
<style>
#MainMenu {visibility: hidden;}
header {visibility: hidden;}
footer {visibility: hidden;}
</style>

"""

st.markdown(hide_streamlit_style, unsafe_allow_html=True)

API_BASE_URL = os.getenv("API_BASE_URL", "https://desertlike-nonrecognized-keagan.ngrok-free.dev")

def count_words(text):
    return len(str(text).split())

def estimate_pages(text, words_per_page=250):
    import math
    return math.ceil(count_words(text) / words_per_page)

def validate_document_length(text, max_words=1000):
    word_count = count_words(text)
    if word_count > max_words:
        return False, f"Document too long! ({word_count} words). Maximum {max_words} words."
    else:
        return True, f"Document valid ({word_count} words)"

def get_text_statistics(text):
    word_count = count_words(text)
    char_count = len(text)
    sentence_count = len(re.findall(r'[.!?]+', text))
    paragraph_count = len([p for p in text.split('\n\n') if p.strip()])
    avg_words_per_sentence = word_count / max(sentence_count, 1)
    complexity_score = min(100, max(0, 100 - (avg_words_per_sentence * 2)))
    return {
        'words': word_count,
        'characters': char_count,
        'sentences': sentence_count,
        'paragraphs': paragraph_count,
        'pages': estimate_pages(text),
        'complexity_score': round(complexity_score, 1),
        'avg_words_per_sentence': round(avg_words_per_sentence, 1)
    }

def extract_text_from_pdf(pdf_file):
    if not PDF_SUPPORT:
        return None, "PDF support not available. Please install PyPDF2 and pdfplumber first."
    try:
        text = ""
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return text.strip(), None
        pdf_file.seek(0)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        if text.strip():
            return text.strip(), None
        else:
            return None, "Cannot extract text from PDF. File may be image-based or corrupted."
    except Exception as e:
        return None, f"Error processing PDF: {str(e)}"

def create_pdf_preview(text, max_chars=500):
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "..."

def validate_pdf_file(uploaded_file):
    if uploaded_file.type != "application/pdf":
        return False, "File must be in PDF format"
    if uploaded_file.size > 10 * 1024 * 1024:
        return False, "PDF file too large! Maximum 10MB"
    return True, "PDF file valid"

def extract_text_from_docx(docx_file):
    if not DOCX_SUPPORT:
        return None, "DOCX support not available. Please install python-docx first."
    try:
        doc = Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        if text.strip():
            return text.strip(), None
        else:
            return None, "No text found in DOCX. File may be empty."
    except Exception as e:
        return None, f"Error processing DOCX: {str(e)}"

def create_docx_preview(text, max_chars=500):
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "..."

def validate_docx_file(uploaded_file):
    valid_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword"
    ]
    if uploaded_file.type not in valid_types:
        if not uploaded_file.name.lower().endswith(('.docx', '.doc')):
            return False, "File must be in DOCX or DOC format"
    if uploaded_file.size > 10 * 1024 * 1024:
        return False, "DOCX file too large! Maximum 10MB"
    return True, "DOCX file valid"

@st.cache_data(ttl=60)
def predict_similarity_api(text1: str, text2: str):
    try:
        payload = {
            "sentence1": text1,
            "sentence2": text2
        }
        response = requests.post(
            f"{API_BASE_URL}/sentence-similarity",
            json=payload,
            timeout=30
        )
        if response.status_code == 200:
            resp = response.json()
            normalized = {
                "similarity": resp.get("probability"),
                "label": resp.get("label", ""),
                "processing_time": resp.get("processing_time", 0)
            }
            return normalized
        else:
            error_detail = response.json().get("detail", "Unknown error")
            return {"error": f"API Error: {error_detail}"}
    except requests.exceptions.Timeout:
        return {"error": "Request timeout - API took too long to respond"}
    except requests.exceptions.RequestException as e:
        return {"error": f"Connection error: {str(e)}"}

def predict_document_api(doc1: str, doc2: str):
    try:
        payload = {
            "text_a": doc1,
            "text_b": doc2
        }
        resp = requests.post(
            f"{API_BASE_URL}/document-similarity",
            json=payload,
            timeout=300
        )
        if resp.status_code == 200:
            data = resp.json()
            analysis = data.get("analysis", {})
            return {
                "similarity": analysis.get("similarity_score") or 0,
                "label": analysis.get("similarity_label") or "Unknown",
                "processing_time": data.get("processing_time") or 0
            }
        else:
            det = resp.json().get("detail", f"HTTP {resp.status_code}")
            return {"error": f"API Error: {det}"}
    except requests.exceptions.Timeout:
        return {"error": "Request timeout - document API took too long to respond"}
    except requests.exceptions.RequestException as e:
        return {"error": f"Connection error: {str(e)}"}

def switch_model_api(target: str, model_filename: Optional[str] = None):
    try:
        params = {"target": target}
        if model_filename:
            params["model_filename"] = model_filename
        resp = requests.post(
            f"{API_BASE_URL}/switch-model",
            params=params,
            timeout=1200
        )
        if resp.status_code == 200:
            return {"success": True, "message": f"Model switched to {target}"}
        else:
            det = resp.json().get("detail", f"HTTP {resp.status_code}")
            return {"success": False, "error": f"API Error: {det}"}
    except requests.exceptions.Timeout:
        return {"success": False, "error": "Request timeout"}
    except requests.exceptions.RequestException as e:
        return {"success": False, "error": f"Connection error: {str(e)}"}

def create_similarity_gauge(similarity, label=""):
    if similarity >= 0.8:
        bar_color = "#f5576c"
    elif similarity >= 0.6:
        bar_color = "#fdcb6e"
    else:
        bar_color = "#00b894"
    
    fig = go.Figure(go.Indicator(
        mode = "gauge+number",
        value = similarity * 100,
        number = {
            'font': {'size': 64, 'color': bar_color, 'family': 'Inter, -apple-system, BlinkMacSystemFont, sans-serif'},
            'suffix': '%',
            'valueformat': '.1f'
        },
        domain = {'x': [0, 1], 'y': [0, 1]},
        gauge = {
            'axis': {
                'range': [0, 100], 
                'tickwidth': 2,
                'tickcolor': '#a0aec0',
                'tickmode': 'array',
                'tickvals': [0, 20, 40, 60, 80, 100],
                'ticktext': ['0', '20', '40', '60', '80', '100'],
                'tickfont': {'size': 14, 'color': '#a0aec0', 'family': 'Inter, sans-serif'}
            },
            'bar': {'color': bar_color, 'thickness': 0.35},
            'bgcolor': 'rgba(255,255,255,0.1)',
            'borderwidth': 2,
            'bordercolor': '#4a5568',
            'steps': [
                {'range': [0, 60], 'color': 'rgba(0, 184, 148, 0.25)'},
                {'range': [60, 80], 'color': 'rgba(253, 203, 110, 0.25)'},
                {'range': [80, 100], 'color': 'rgba(245, 87, 108, 0.25)'}
            ],
            'threshold': {
                'line': {'color': bar_color, 'width': 5},
                'thickness': 0.8,
                'value': similarity * 100
            }
        }
    ))
    fig.update_layout(
        height=400, 
        font={'color': "#ffffff", 'family': "Inter, sans-serif"},
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)',
        margin=dict(l=40, r=40, t=40, b=40)
    )
    return fig

initialize_auth_state()

st.markdown("""
<div class="main-header">
    <h1>Semantic Textual Similarity Application</h1>
    <h4>Universitas Mikroskil</h4>
    <p>Chrisandy | Fahim | Sandy</p>
    <p>2025 / 2026</p>
</div>
""", unsafe_allow_html=True)

st.sidebar.markdown("## Navigation Menu")
analysis_type = st.sidebar.radio(
    "Select analysis type:",
    ["📄 Text Similarity", "📁 Document Similarity"],
    index=0
)

with st.sidebar:
    c1, c2 = st.columns([1, 12])
    with c1:
        st.image("images/model.png", width=24)
    with c2:
        st.markdown("Base Model")

st.sidebar.markdown(
    """
    <div style="margin:0; line-height:1.2">
      <strong>Embedding:</strong> BERT-Based-Uncased<br>
      <strong>Neural Model:</strong> BiLSTM + Attention<br>
      <strong>Dataset:</strong> Quora Question Pairs
    </div>
    """,
    unsafe_allow_html=True
)

st.sidebar.markdown("---")

if 'selected_model' not in st.session_state:
    st.session_state.selected_model = 'manual'

@st.cache_data(ttl=60)
def get_available_models():
    try:
        resp = requests.get(f"{API_BASE_URL}/list-models", timeout=10)
        resp.raise_for_status()
        return resp.json()
    except Exception as e:
        return None

if is_logged_in():
    with st.sidebar:
        current_model_display = "Base Model" if st.session_state.selected_model == 'manual' else "Retrain Model"
        st.markdown(f"**Current Model:** {current_model_display}")
        
        st.markdown("### 🔄 Switch Model")
        
        models_data = get_available_models()
        
        if models_data:
            colab_models = models_data.get("available_colab_models", [])
            
            if colab_models:
                model_option = st.selectbox(
                    "Select model:",
                    options=["manual"] + colab_models,
                    format_func=lambda x: "Base Model (manual)" if x == "manual" else f"Retrain: {x}",
                    index=0 if st.session_state.selected_model == 'manual' else min(1, len(colab_models) + 1),
                    key="model_selector"
                )
                
                if st.button("Switch Model", width='stretch', type="secondary"):
                    with st.spinner(f"Switching to model..."):
                        if model_option == "manual":
                            result = switch_model_api("manual")
                        else:
                            result = switch_model_api("colab", model_option)
                        
                        if result.get("success"):
                            st.session_state.selected_model = model_option
                            st.cache_data.clear()
                            st.success(f"✅ Model switched")
                            st.rerun()
                        else:
                            st.error(f"❌ {result.get('error', 'Failed to switch')}")
            else:
                # No colab models, fallback ke manual only
                st.info("Only base model available")
                if st.button("Use Base Model", width='stretch', type="secondary"):
                    st.session_state.selected_model = 'manual'
                    st.success("✅ Using base model")
        else:
            st.warning("Could not fetch models")

st.sidebar.markdown("---")

with st.sidebar:
    if is_logged_in():
        if st.button("Admin Panel", key="admin_sidebar_logged", width='stretch', type="primary"):
            st.switch_page("pages/admin_panel.py")
    else:
        if st.button("Login Admin", key="admin_sidebar_guest", width='stretch'):
            st.switch_page("pages/admin_panel.py")

# Initialize history in session state
if 'text_similarity_history' not in st.session_state:
    st.session_state.text_similarity_history = []
if 'doc_similarity_history' not in st.session_state:
    st.session_state.doc_similarity_history = []

if analysis_type == "📄 Text Similarity":
    st.markdown("## Text Similarity Analysis")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("### 📝 Text 1")
        text1 = st.text_area("Enter first text (max 30 words):", height=200, key="text1", 
                            placeholder="Paste or type first text here...")
        if text1:
            word_count1 = len(text1.split())
            if word_count1 == 0:
                st.error("Text cannot be empty")
                st.session_state.text1_valid = False
            elif word_count1 > 30:
                st.error(f"Text too long: {word_count1} words (max 30 words)")
                st.session_state.text1_valid = False
            else:
                st.success(f"Text valid, {word_count1} words")
                st.session_state.text1_valid = True
        else:
            st.session_state.text1_valid = False
    with col2:
        st.markdown("### 📝 Text 2")
        text2 = st.text_area("Enter second text (max 30 words):", height=200, key="text2",
                            placeholder="Paste or type second text here...")
        if text2:
            word_count2 = len(text2.split())
            if word_count2 == 0:
                st.error("Text cannot be empty")
                st.session_state.text2_valid = False
            elif word_count2 > 30:
                st.error(f"Text too long: {word_count2} words (max 30 words)")
                st.session_state.text2_valid = False
            else:
                st.success(f"Text valid, {word_count2} words")
                st.session_state.text2_valid = True
        else:
            st.session_state.text2_valid = False
    both_texts_valid = st.session_state.get('text1_valid', False) and st.session_state.get('text2_valid', False)
    text_button_disabled = not both_texts_valid
    if not both_texts_valid:
        if not st.session_state.get('text1_valid', False) and not st.session_state.get('text2_valid', False):
            st.info("Enter both valid texts to start analysis (max 30 words per text)")
        elif not st.session_state.get('text1_valid', False):
            st.warning("Text 1 is not valid or empty")
        elif not st.session_state.get('text2_valid', False):
            st.warning("Text 2 is not valid or empty")
    col_center = st.columns([1, 2, 1])[1]
    with col_center:
        analyze_button = st.button(
            "Check Similarity", 
            type="primary",
            disabled=text_button_disabled,
            width='stretch'
        )
        if analyze_button:
            if not text1 or not text2:
                st.error("Please enter both valid texts!")
            else:
                with st.spinner():
                    result = predict_similarity_api(text1, text2)
                err = result.get("error") or result.get("detail")
                if err:
                    st.error(f"{err}")
                else:
                    similarity = result['similarity']
                    label = result['label']
                    processing_time = result['processing_time']
                    # Add to history
                    import datetime
                    st.session_state.text_similarity_history.insert(0, {
                        'text1_preview': text1[:50] + '...' if len(text1) > 50 else text1,
                        'text2_preview': text2[:50] + '...' if len(text2) > 50 else text2,
                        'similarity': similarity,
                        'label': label,
                        'processing_time': processing_time,
                        'timestamp': datetime.datetime.now().strftime('%H:%M:%S')
                    })
                    # Keep only last 10 entries
                    st.session_state.text_similarity_history = st.session_state.text_similarity_history[:10]
                    st.markdown("---")
                    st.markdown("<h2 style='text-align: center;'>Similarity Analysis Results</h2>", unsafe_allow_html=True)
                    
                    perf_cols = st.columns(3)
                    with perf_cols[0]:
                        st.markdown(f"""
                        <div style='background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
                                    padding: 1.5rem; border-radius: 15px; border-left: 4px solid #667eea;'>
                            <div style='color: #667eea; font-size: 0.85rem; font-weight: 600; margin-bottom: 0.5rem;'>
                                SIMILARITY SCORE
                            </div>
                            <div style='font-size: 2rem; font-weight: 800; color: #ffffff;'>
                                {similarity:.4f}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    with perf_cols[1]:
                        st.markdown(f"""
                        <div style='background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
                                    padding: 1.5rem; border-radius: 15px; border-left: 4px solid #764ba2;'>
                            <div style='color: #764ba2; font-size: 0.85rem; font-weight: 600; margin-bottom: 0.5rem;'>
                                LABEL
                            </div>
                            <div style='font-size: 1.5rem; font-weight: 800; color: #ffffff;'>
                                {label}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    with perf_cols[2]:
                        st.markdown(f"""
                        <div style='background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
                                    padding: 1.5rem; border-radius: 15px; border-left: 4px solid #28a745;'>
                            <div style='color: #28a745; font-size: 0.85rem; font-weight: 600; margin-bottom: 0.5rem;'>
                                PROCESSING TIME
                            </div>
                            <div style='font-size: 2rem; font-weight: 800; color: #ffffff;'>
                                {processing_time:.2f}s
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    gauge_fig = create_similarity_gauge(similarity, label)
                    st.plotly_chart(gauge_fig, use_container_width=True)
    
    # Show Text Similarity History
    if st.session_state.text_similarity_history:
        st.markdown("---")
        st.markdown("### 📜 Analysis History")
        
        
        for i, entry in enumerate(st.session_state.text_similarity_history):
            similarity_color = "#28a745" if entry['similarity'] < 0.6 else "#ffc107" if entry['similarity'] < 0.8 else "#dc3545"
            with st.expander(f"🕐 {entry['timestamp']} | Score: {entry['similarity']:.4f} | {entry['label']}", expanded=(i==0)):
                col_h1, col_h2 = st.columns(2)
                with col_h1:
                    st.markdown(f"**Text 1:** {entry['text1_preview']}")
                with col_h2:
                    st.markdown(f"**Text 2:** {entry['text2_preview']}")
                st.markdown(f"⏱️ Processing time: **{entry['processing_time']:.2f}s**")
        
        if st.button("🗑️ Clear History", key="clear_text_history"):
            st.session_state.text_similarity_history = []
            st.rerun()

elif analysis_type == "📁 Document Similarity":
    st.markdown("## Document Similarity Analysis")
    supported_formats = []
    if PDF_SUPPORT:
        supported_formats.append("PDF")
    if DOCX_SUPPORT:
        supported_formats.append("DOCX")
    supported_formats.append("TXT")
    if not PDF_SUPPORT and not DOCX_SUPPORT:
        st.warning("**PDF & DOCX support not available.** Install PyPDF2, pdfplumber, and python-docx for full support.")
        st.info("Upload exactly 2 TXT documents to compare (max 1000 words per document)")
    else:
        st.info(f"Upload exactly 2 documents ({', '.join(supported_formats)}) to compare (max 1000 words per document)")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("### 📄 Document 1")
        file_types = ["txt"]
        if PDF_SUPPORT:
            file_types.append("pdf")
        if DOCX_SUPPORT:
            file_types.extend(["docx", "doc"])
        uploaded_file1 = st.file_uploader(
            f"Upload first document ({', '.join([t.upper() for t in file_types])})", 
            type=file_types, 
            key="file1"
        )
        doc1_content = ""
        if uploaded_file1:
            if uploaded_file1.name.lower().endswith(('.docx', '.doc')):
                st.session_state.doc1_valid = False
                if DOCX_SUPPORT:
                    is_valid_docx, docx_msg = validate_docx_file(uploaded_file1)
                    if is_valid_docx:
                        with st.spinner("📖 Extracting text from DOCX..."):
                            text, error = extract_text_from_docx(uploaded_file1)
                        if error:
                            st.error(error)
                            st.session_state.doc1_valid = False
                        else:
                            doc1_content = text
                            is_valid1, msg1 = validate_document_length(doc1_content)
                            if is_valid1:
                                st.success("DOCX processed successfully!")
                                st.success(msg1)
                                st.session_state.doc1_content = doc1_content
                                with st.expander("Preview Document 1 (from DOCX)"):
                                    st.text_area("Content:", create_docx_preview(doc1_content), 
                                               height=150, disabled=True, key=f"preview1_docx_{uploaded_file1.name}_{uploaded_file1.size}")
                                st.session_state.doc1_valid = True
                            else:
                                st.error(msg1)
                                doc1_content = ""
                                st.session_state.doc1_valid = False
                    else:
                        st.error(docx_msg)
                        st.session_state.doc1_valid = False
                else:
                    st.error("DOCX support not available")
                    st.session_state.doc1_valid = False
            elif uploaded_file1.type == "application/pdf":
                st.session_state.doc1_valid = False
                if PDF_SUPPORT:
                    is_valid_pdf, pdf_msg = validate_pdf_file(uploaded_file1)
                    if is_valid_pdf:
                        with st.spinner("📖 Extracting text from PDF..."):
                            text, error = extract_text_from_pdf(uploaded_file1)
                        if error:
                            st.error(error)
                            st.session_state.doc1_valid = False
                        else:
                            doc1_content = text
                            is_valid1, msg1 = validate_document_length(doc1_content)
                            if is_valid1:
                                st.success("PDF processed successfully!")
                                st.success(msg1)
                                st.session_state.doc1_content = doc1_content
                                with st.expander("Preview Document 1 (from PDF)"):
                                    st.text_area("Content:", create_pdf_preview(doc1_content), 
                                               height=150, disabled=True, key=f"preview1_pdf_{uploaded_file1.name}_{uploaded_file1.size}")
                                st.session_state.doc1_valid = True
                            else:
                                st.error(msg1)
                                doc1_content = ""
                                st.session_state.doc1_valid = False
                    else:
                        st.error(pdf_msg)
                        st.session_state.doc1_valid = False
                else:
                    st.error("PDF support not available")
                    st.session_state.doc1_valid = False
            else:
                st.session_state.doc1_valid = False
                doc1_content = uploaded_file1.read().decode("utf-8")
                is_valid1, msg1 = validate_document_length(doc1_content)
                if is_valid1:
                    st.success("TXT file loaded successfully!")
                    st.success(msg1)
                    st.session_state.doc1_content = doc1_content
                    with st.expander("Preview Document 1"):
                        st.text_area("Content:", doc1_content[:500] + "..." if len(doc1_content) > 500 else doc1_content, 
                                   height=150, disabled=True, key=f"preview1_txt_{uploaded_file1.name}_{uploaded_file1.size}")
                    st.session_state.doc1_valid = True
                else:
                    st.error(msg1)
                    doc1_content = ""
                    st.session_state.doc1_valid = False
        else:
            st.session_state.doc1_valid = False
    with col2:
        st.markdown("### 📄 Document 2")
        uploaded_file2 = st.file_uploader(
            f"Upload second document ({', '.join([t.upper() for t in file_types])})", 
            type=file_types, 
            key="file2"
        )
        doc2_content = ""
        if uploaded_file2:
            if uploaded_file2.name.lower().endswith(('.docx', '.doc')):
                st.session_state.doc2_valid = False
                if DOCX_SUPPORT:
                    is_valid_docx, docx_msg = validate_docx_file(uploaded_file2)
                    if is_valid_docx:
                        with st.spinner("📖 Extracting text from DOCX..."):
                            text, error = extract_text_from_docx(uploaded_file2)
                        if error:
                            st.error(error)
                            st.session_state.doc2_valid = False
                        else:
                            doc2_content = text
                            is_valid2, msg2 = validate_document_length(doc2_content)
                            if is_valid2:
                                st.success("DOCX processed successfully!")
                                st.success(msg2)
                                st.session_state.doc2_content = doc2_content
                                with st.expander("Preview Document 2 (from DOCX)"):
                                    st.text_area("Content:", create_docx_preview(doc2_content), 
                                               height=150, disabled=True, key=f"preview2_docx_{uploaded_file2.name}_{uploaded_file2.size}")
                                st.session_state.doc2_valid = True
                            else:
                                st.error(msg2)
                                doc2_content = ""
                                st.session_state.doc2_valid = False
                    else:
                        st.error(docx_msg)
                        st.session_state.doc2_valid = False
                else:
                    st.error("DOCX support not available")
                    st.session_state.doc2_valid = False
            elif uploaded_file2.type == "application/pdf":
                st.session_state.doc2_valid = False
                if PDF_SUPPORT:
                    is_valid_pdf, pdf_msg = validate_pdf_file(uploaded_file2)
                    if is_valid_pdf:
                        with st.spinner("📖 Extracting text from PDF..."):
                            text, error = extract_text_from_pdf(uploaded_file2)
                        if error:
                            st.error(error)
                            st.session_state.doc2_valid = False
                        else:
                            doc2_content = text
                            is_valid2, msg2 = validate_document_length(doc2_content)
                            if is_valid2:
                                st.success("PDF processed successfully!")
                                st.success(msg2)
                                st.session_state.doc2_content = doc2_content
                                with st.expander("Preview Document 2 (from PDF)"):
                                    st.text_area("Content:", create_pdf_preview(doc2_content), 
                                               height=150, disabled=True, key=f"preview2_pdf_{uploaded_file2.name}_{uploaded_file2.size}")
                                st.session_state.doc2_valid = True
                            else:
                                st.error(msg2)
                                doc2_content = ""
                                st.session_state.doc2_valid = False
                    else:
                        st.error(pdf_msg)
                        st.session_state.doc2_valid = False
                else:
                    st.error("PDF support not available")
                    st.session_state.doc2_valid = False
            else:
                st.session_state.doc2_valid = False
                doc2_content = uploaded_file2.read().decode("utf-8")
                is_valid2, msg2 = validate_document_length(doc2_content)
                if is_valid2:
                    st.success("TXT file loaded successfully!")
                    st.success(msg2)
                    st.session_state.doc2_content = doc2_content
                    with st.expander("Preview Document 2"):
                        st.text_area("Content:", doc2_content[:500] + "..." if len(doc2_content) > 500 else doc2_content, 
                                   height=150, disabled=True, key=f"preview2_txt_{uploaded_file2.name}_{uploaded_file2.size}")
                    st.session_state.doc2_valid = True
                else:
                    st.error(msg2)
                    doc2_content = ""
                    st.session_state.doc2_valid = False
        else:
            st.session_state.doc2_valid = False
    both_docs_valid = st.session_state.get('doc1_valid', False) and st.session_state.get('doc2_valid', False)
    button_disabled = not both_docs_valid
    if not both_docs_valid:
        if not st.session_state.get('doc1_valid', False) and not st.session_state.get('doc2_valid', False):
            st.info("Upload both valid documents to start analysis")
        elif not st.session_state.get('doc1_valid', False):
            st.warning("Document 1 is not valid or not uploaded")
        elif not st.session_state.get('doc2_valid', False):
            st.warning("Document 2 is not valid or not uploaded")
    col_center = st.columns([1, 2, 1])[1]
    with col_center:
        doc_analyze_button = st.button("Check Similarity", type="primary", disabled=button_disabled, width='stretch')
        if doc_analyze_button:
            # Get content from session state
            doc1_content = st.session_state.get('doc1_content', '')
            doc2_content = st.session_state.get('doc2_content', '')
            if not doc1_content or not doc2_content:
                st.error("Please upload both valid documents!")
            else:
                with st.spinner("Analyzing documents..."):
                        result = predict_document_api(doc1_content, doc2_content)
                if "error" in result:
                    st.error(f"{result['error']}")
                else:
                    similarity = result["similarity"] or 0
                    label = result["label"] or "Unknown"
                    processing_time = result["processing_time"] or 0
                    st.markdown("---")
                    st.markdown("<h2 style='text-align: center;'>Document Analysis Results</h2>", unsafe_allow_html=True)
                    
                    perf_cols = st.columns(3)
                    with perf_cols[0]:
                        st.markdown(f"""
                        <div style='background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
                                    padding: 1.5rem; border-radius: 15px; border-left: 4px solid #667eea;'>
                            <div style='color: #667eea; font-size: 0.85rem; font-weight: 600; margin-bottom: 0.5rem;'>
                                SIMILARITY SCORE
                            </div>
                            <div style='font-size: 2rem; font-weight: 800; color: #ffffff;'>
                                {similarity:.4f}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    with perf_cols[1]:
                        st.markdown(f"""
                        <div style='background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
                                    padding: 1.5rem; border-radius: 15px; border-left: 4px solid #764ba2;'>
                            <div style='color: #764ba2; font-size: 0.85rem; font-weight: 600; margin-bottom: 0.5rem;'>
                                LABEL
                            </div>
                            <div style='font-size: 1.5rem; font-weight: 800; color: #ffffff;'>
                                {label}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    with perf_cols[2]:
                        st.markdown(f"""
                        <div style='background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
                                    padding: 1.5rem; border-radius: 15px; border-left: 4px solid #28a745;'>
                            <div style='color: #28a745; font-size: 0.85rem; font-weight: 600; margin-bottom: 0.5rem;'>
                                PROCESSING TIME
                            </div>
                            <div style='font-size: 2rem; font-weight: 800; color: #ffffff;'>
                                {processing_time:.2f}s
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    gauge_fig = create_similarity_gauge(similarity, label)
                    st.plotly_chart(gauge_fig, use_container_width=True)
                    
                    # Add to document history
                    import datetime
                    st.session_state.doc_similarity_history.insert(0, {
                        'doc1_name': uploaded_file1.name if uploaded_file1 else 'Unknown',
                        'doc2_name': uploaded_file2.name if uploaded_file2 else 'Unknown',
                        'similarity': similarity,
                        'label': label,
                        'processing_time': processing_time,
                        'timestamp': datetime.datetime.now().strftime('%H:%M:%S')
                    })
                    # Keep only last 10 entries
                    st.session_state.doc_similarity_history = st.session_state.doc_similarity_history[:10]
    
    # Show Document Similarity History
    if st.session_state.doc_similarity_history:
        st.markdown("---")
        st.markdown("### 📜 Analysis History")
        
        
        for i, entry in enumerate(st.session_state.doc_similarity_history):
            similarity_color = "#28a745" if entry['similarity'] < 0.6 else "#ffc107" if entry['similarity'] < 0.8 else "#dc3545"
            with st.expander(f"🕐 {entry['timestamp']} | Score: {entry['similarity']:.4f} | {entry['label']}", expanded=(i==0)):
                st.markdown(f"**Document 1:** {entry['doc1_name']}")
                st.markdown(f"**Document 2:** {entry['doc2_name']}")
                st.markdown(f"⏱️ Processing time: **{entry['processing_time']:.2f}s**")
        
        if st.button("🗑️ Clear History", key="clear_doc_history"):
            st.session_state.doc_similarity_history = []
            st.rerun()

st.markdown("---")
